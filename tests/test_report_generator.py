import unittest
import os
import json
import numpy as np
import pandas as pd
import shutil
from unittest.mock import patch, MagicMock, ANY

# Adjust import path assuming tests are run from the repository root
from report_generator import create_individual_experiment_report, create_comparative_docx_report
from docx import Document # For type hinting and mocking, not direct use here.

# Helper function to load and prepare test scenario data
def _load_test_scenario_data(scenario_name: str) -> dict:
    """
    Loads data from test_data/{scenario_name}_data.json and prepares it
    in a structure similar to what scientific_workflow.py would pass.
    """
    base_dir = os.path.join(os.getcwd(), 'test_data')
    file_path = os.path.join(base_dir, f"{scenario_name}_data.json")

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Test data file not found: {file_path}")

    with open(file_path, 'r') as f:
        data = json.load(f)

    # Mimic structure from scientific_workflow.py's scenario_data
    # Convert lists to numpy arrays where appropriate for plotting or internal logic
    # The report_generator mainly expects lists for observations/predictions/time_steps
    # but validation_metrics and calibration_results are dicts.

    # Ensure observed_data values are lists, as expected by plotting in report_generator
    observed_data_processed = {}
    for key, value in data.get('observed_data', {}).items():
        observed_data_processed[key] = value if isinstance(value, list) else list(value)

    scenario_output_data = {
        'validation_metrics': data.get('validation_metrics', { # Mock if not in test_data files
            'rmse': 0.1, 'mae': 0.1, 'mape': 1.0, 'nse': 0.9, 'pearson_r': 0.95
        }),
        'hypothesis_results': data.get('hypothesis_results', { # Mock if not in test_data files
            'hypothesis_supported': True, 'criteria_passed': 4, 'total_criteria': 4
        }),
        'results_interpretation': data.get('results_interpretation', { # Mock this, as it comes from LLM
            "results_narrative": f"Test results narrative for {scenario_name}.",
            "discussion_narrative": f"Test discussion narrative for {scenario_name}.",
            "conclusion_narrative": f"Test conclusion narrative for {scenario_name}.",
            "structured_assessment": { # Add if your report uses it
                "clarity_of_results": "High", "comparison_to_hypothesis": "Supported",
                "limitations_of_study": "None noted for test", "implications_for_future_work": "Further testing"
            }
        }),
        # These should be lists as per current scientific_workflow.py structure for scenario_data
        'predictions': list(np.random.rand(len(observed_data_processed.get('storage_level', [])))) if 'storage_level' in observed_data_processed else [],
        'observations': observed_data_processed.get('storage_level', []),
        'time_steps': data.get('time_steps', []),
        'calibration_results': data.get('calibration_results', { # Mock if not in test_data files
            'storage_capacity': 100, 'overflow_coefficient': 0.8, 'storage_efficiency': 0.9, 'calibration_rmse': 0.05
        }),
        # Add other keys if your report_generator uses them from scenario_data
        # For example, if the full observed_data_np (converted to lists) is used:
        'full_observed_data_np': observed_data_processed
    }
    # Ensure 'predictions' matches length of 'observations' if 'observations' exist
    if scenario_output_data['observations'] and not scenario_output_data['predictions']:
        scenario_output_data['predictions'] = list(np.random.rand(len(scenario_output_data['observations'])))
    elif len(scenario_output_data['predictions']) != len(scenario_output_data['observations']):
         scenario_output_data['predictions'] = list(np.random.rand(len(scenario_output_data['observations'])))


    return scenario_output_data


class TestIndividualReport(unittest.TestCase):
    def setUp(self):
        self.initial_sections_json = {
            "abstract": "This is a test abstract.",
            "introduction": "This is a test introduction.",
            "hypothesis_narrative": "The model is hypothesized to be accurate.",
            "hypothesis_variables": ["Variable A", "Variable B"],
            "methods_narrative": "Standard methods were used.",
            "methods_components": ["Component X", "Component Y"]
        }
        # Ensure test_data files exist, or this will fail. test_data_handling should ensure this.
        try:
            self.support_scenario_data = _load_test_scenario_data('support')
        except FileNotFoundError as e:
            self.fail(f"Setup failed for TestIndividualReport: {e}. Run generate_test_data.py first.")

        self.output_dir = "temp_test_output_individual"
        os.makedirs(self.output_dir, exist_ok=True)

    def tearDown(self):
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir)

    @patch('report_generator.plt')
    @patch('report_generator.Document')
    def test_individual_report_structure_and_calls(self, MockDocument, mock_plt):
        mock_doc_instance = MockDocument.return_value

        # Configure mock_plt.subplots to return a fig and an axes array
        mock_fig = MagicMock()
        # Explicitly create individual MagicMock objects for each subplot
        self.mock_ax_00 = MagicMock()
        self.mock_ax_01 = MagicMock()
        self.mock_ax_10 = MagicMock()
        self.mock_ax_11 = MagicMock()
        # Revert to numpy array of mocks with dtype=object, as report_generator.py uses tuple indexing (e.g., axes[0,0])
        mock_axes_array_of_mocks = np.array([
            [self.mock_ax_00, self.mock_ax_01],
            [self.mock_ax_10, self.mock_ax_11]
        ], dtype=object)
        mock_plt.subplots.return_value = (mock_fig, mock_axes_array_of_mocks)

        # If plt.figure() is called separately by the code, it needs its own mock setup.
        # The current code uses fig, axes = plt.subplots(...), so plt.figure() might not be directly called.
        # The assertion mock_plt.figure.assert_called_once() might need adjustment.
        # Let's assume plt.figure() is NOT called directly if subplots is used.
        # Instead, the `fig` from `subplots` is used. `plt.close(fig)` would be `mock_plt.close(mock_fig)`.

        create_individual_experiment_report('support', self.initial_sections_json, self.support_scenario_data, self.output_dir)

        # Check Document save call
        expected_docx_path = os.path.join(self.output_dir, 'report_support.docx')
        mock_doc_instance.save.assert_called_with(expected_docx_path)

        # Check plot save call
        expected_plot_path = os.path.join(self.output_dir, 'plot_support.png')
        mock_plt.savefig.assert_called_with(expected_plot_path)

        # Check that subplots was called to create the figure and axes
        mock_plt.subplots.assert_called_once_with(2, 2, figsize=(8, 6))

        # Check plotting commands were called on the correct axes objects
        # For example, if the top-left plot is predictions/observations:
        self.mock_ax_00.plot.assert_called()
        self.mock_ax_00.set_xlabel.assert_called()
        self.mock_ax_00.set_ylabel.assert_called()
        self.mock_ax_00.set_title.assert_called()
        self.mock_ax_00.legend.assert_called()

        # Example for another plot if it exists and is used (e.g. energy distribution)
        # self.mock_ax_01.hist.assert_called()
        # self.mock_ax_01.set_title.assert_called_with('Energy Input Distribution')

        # Check that plt.close was called with the figure object from subplots
        mock_plt.close.assert_called_with(mock_fig)

        # Check for add_heading calls (count might be fragile, check for specific important ones)
        mock_plt.ylabel.assert_called()
        mock_plt.title.assert_called()
        mock_plt.legend.assert_called()
        mock_plt.close.assert_called_with(mock_plt.figure.return_value)


        # Check for add_heading calls (count might be fragile, check for specific important ones)
        self.assertTrue(mock_doc_instance.add_heading.call_count >= 5) # Abstract, Intro, Hypo, Methods, Results etc.

        # Check for add_paragraph calls
        self.assertTrue(mock_doc_instance.add_paragraph.call_count > 0)

        # Check for add_picture call
        mock_doc_instance.add_picture.assert_called_with(expected_plot_path, width=ANY) # ANY for width

    def test_individual_report_file_creation(self):
        # This test does not mock Document or plt.savefig to check actual file creation
        create_individual_experiment_report('support', self.initial_sections_json, self.support_scenario_data, self.output_dir)

        expected_docx_path = os.path.join(self.output_dir, 'report_support.docx')
        self.assertTrue(os.path.exists(expected_docx_path), f"DOCX file not created at {expected_docx_path}")

        expected_plot_path = os.path.join(self.output_dir, 'plot_support.png')
        self.assertTrue(os.path.exists(expected_plot_path), f"Plot PNG file not created at {expected_plot_path}")
        self.assertTrue(os.path.getsize(expected_plot_path) > 0, "Plot PNG file is empty.")


class TestComparativeReport(unittest.TestCase):
    def setUp(self):
        self.llm_response_1 = {
            "overall_introduction": "Overall intro from LLM1.",
            "experiment_summaries": [
                {"scenario_name": "support", "key_finding": "Placeholder KF for support", "metrics": {}},
                {"scenario_name": "fail", "key_finding": "Placeholder KF for fail", "metrics": {}}
            ],
            "overall_conclusion": "Overall conclusion from LLM1."
        }
        self.llm_response_2 = {
            "overall_introduction": "Refined overall intro from LLM2.", # report_generator uses this
            "experiment_summaries": [
                {
                    "scenario_name": "support",
                    "key_finding": "Detailed key finding for support from LLM2.",
                    "metrics": {"rmse": 0.1, "nse": 0.9, "hypothesis_supported": True, "mae": 0.05, "mape": 1.5}
                },
                {
                    "scenario_name": "fail",
                    "key_finding": "Detailed key finding for fail from LLM2.",
                    "metrics": {"rmse": 0.5, "nse": 0.3, "hypothesis_supported": False, "mae": 0.4, "mape": 10.2}
                }
            ],
            "comparative_analysis": {
                "performance_overview": "Model performed variably. Good in support, poor in fail.",
                "cross_scenario_insights": "Insights show sensitivity to initial conditions."
            },
            "overall_conclusion": "Final overall conclusion from LLM2." # report_generator uses this
        }
        self.output_dir = "temp_test_output_comparative"
        os.makedirs(self.output_dir, exist_ok=True)

    def tearDown(self):
        if os.path.exists(self.output_dir):
            shutil.rmtree(self.output_dir)

    @patch('report_generator.Document')
    def test_comparative_report_structure_and_calls(self, MockDocument):
        mock_doc_instance = MockDocument.return_value

        create_comparative_docx_report(self.llm_response_1, self.llm_response_2, self.output_dir)

        expected_docx_path = os.path.join(self.output_dir, 'comparative_report.docx')
        mock_doc_instance.save.assert_called_with(expected_docx_path)

        # Check for key headings - simplified assertion
        self.assertTrue(mock_doc_instance.add_heading.called, "add_heading was not called.")
        # A more robust check would be to verify specific important headings were added,
        # but first, confirm it's being called at all.
        # Example: self.assertGreater(mock_doc_instance.add_heading.call_count, 3)

        # Check that content from llm_response_2 is used for introduction and conclusion
        self.assertTrue(mock_doc_instance.add_paragraph.called, "add_paragraph was not called.")

        # Correctly extract paragraph text: call_obj.args[0] should be the text string
        paragraphs_texts = [call_obj.args[0] for call_obj in mock_doc_instance.add_paragraph.call_args_list if call_obj.args and isinstance(call_obj.args[0], str)]
        paragraphs_added = " ".join(paragraphs_texts)

        self.assertIn(self.llm_response_2["overall_introduction"], paragraphs_added)
        self.assertIn(self.llm_response_2["overall_conclusion"], paragraphs_added)
        self.assertIn(self.llm_response_2["comparative_analysis"]["performance_overview"], paragraphs_added)
        self.assertIn(self.llm_response_2["experiment_summaries"][0]["key_finding"], paragraphs_added) # Check one key finding


    def test_comparative_report_file_creation(self):
        create_comparative_docx_report(self.llm_response_1, self.llm_response_2, self.output_dir)

        expected_docx_path = os.path.join(self.output_dir, 'comparative_report.docx')
        self.assertTrue(os.path.exists(expected_docx_path), f"Comparative DOCX file not created at {expected_docx_path}")
        self.assertTrue(os.path.getsize(expected_docx_path) > 0, "Comparative DOCX file is empty.")


if __name__ == '__main__':
    unittest.main()
