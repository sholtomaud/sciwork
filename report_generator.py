import os
import pandas as pd
import numpy as np # Added import
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt

def add_heading(doc, text, level, size=None):
    """Adds a styled heading to the document."""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = 'Calibri'
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(14) if level == 1 else Pt(12)
    run.bold = True

def add_paragraph(doc, text, style=None):
    """Adds a styled paragraph to the document."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

def create_individual_experiment_report(scenario_name, initial_sections_json, scenario_specific_data, output_subdir):
    """
    Generates an individual scientific report in DOCX format for a specific scenario.
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    title_run = title.add_run(f'Validation of an Odum Energy Systems Model - {scenario_name.title()} Scenario')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    add_heading(doc, 'Abstract', level=1)
    add_paragraph(doc, initial_sections_json.get('abstract', 'N/A'))

    # Introduction
    add_heading(doc, 'Introduction', level=1)
    add_paragraph(doc, initial_sections_json.get('introduction', 'N/A'))

    # Hypothesis Section from initial_sections_json
    add_heading(doc, 'Hypothesis', level=1)
    hypothesis_data = initial_sections_json.get('hypothesis', {})
    
    add_heading(doc, 'Hypothesis Statement', level=2, size=12)
    add_paragraph(doc, hypothesis_data.get('narrative_statement', 'N/A'))

    iv_data = hypothesis_data.get('independent_variable', {})
    dv_data = hypothesis_data.get('dependent_variable', {})

    add_heading(doc, 'Variables', level=2, size=12)
    add_paragraph(doc, 'Independent Variable:', style='List Bullet')
    add_paragraph(doc, f"Name: {iv_data.get('name', 'N/A')}", style='List Bullet 2')
    add_paragraph(doc, f"Description: {iv_data.get('description', 'N/A')}", style='List Bullet 2')
    add_paragraph(doc, f"Unit: {iv_data.get('unit', 'N/A')}", style='List Bullet 2')

    add_paragraph(doc, 'Dependent Variable:', style='List Bullet')
    add_paragraph(doc, f"Name: {dv_data.get('name', 'N/A')}", style='List Bullet 2')
    add_paragraph(doc, f"Description: {dv_data.get('description', 'N/A')}", style='List Bullet 2')
    add_paragraph(doc, f"Unit: {dv_data.get('unit', 'N/A')}", style='List Bullet 2')

    # Methods Section from initial_sections_json
    add_heading(doc, 'Methods', level=1)
    methods_data = initial_sections_json.get('methods', {})
    add_paragraph(doc, methods_data.get('narrative', 'N/A'))

    key_components = methods_data.get('key_components', [])
    if key_components:
        add_heading(doc, 'Summary of Methodological Components', level=2, size=12)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'Purpose in this Study'
        for component in key_components:
            row_cells = table.add_row().cells
            row_cells[0].text = component.get('component_name', 'N/A')
            row_cells[1].text = component.get('purpose_in_study', 'N/A')

    # Scenario Specific Data Section
    # doc.add_page_break() # Optional: decide if each scenario report part needs a page break from methods
    add_heading(doc, f'Scenario Analysis: {scenario_name.title()}', level=1)

    # Plotting
    fig, axes = plt.subplots(2, 2, figsize=(8, 6))
    fig.suptitle(f'Model Validation Results - {scenario_name.title()}', fontsize=14)

    time_steps = scenario_specific_data.get('time_steps', [])
    observations = scenario_specific_data.get('observations', [])
    predictions = scenario_specific_data.get('predictions', [])

    # Ensure observations and predictions are not empty before plotting/calculations
    if not time_steps or not isinstance(observations, list) or not isinstance(predictions, list) or len(observations) == 0 or len(predictions) == 0:
        add_paragraph(doc, "Warning: Observation or prediction data is missing or empty, skipping plot generation.")
    else:
        axes[0,0].plot(time_steps, observations, 'b-', label='Observed', alpha=0.7)
        axes[0,0].plot(time_steps, predictions, 'r--', label='Predicted', alpha=0.7)
        axes[0,0].set_title('Time Series Comparison')
        axes[0,0].legend()
        axes[0,0].grid(True, alpha=0.3)

        axes[0,1].scatter(observations, predictions, alpha=0.6)
        # Handle potential empty sequences for min/max
        min_val = min(min(observations, default=0), min(predictions, default=0))
        max_val = max(max(observations, default=0), max(predictions, default=0))
        if min_val == max_val: # Avoid issues with plotting if all values are the same
            min_val -=1
            max_val +=1
        axes[0,1].plot([min_val, max_val], [min_val, max_val], 'r--', label='1:1 line')
        axes[0,1].set_title('Predicted vs Observed')
        axes[0,1].legend()
        axes[0,1].grid(True, alpha=0.3)

        # Ensure numpy is used for array operations if data is present
        residuals = np.array(observations) - np.array(predictions)
        axes[1,0].scatter(predictions, residuals, alpha=0.6)
        axes[1,0].axhline(y=0, color='r', linestyle='--')
        axes[1,0].set_title('Residual Plot')
        axes[1,0].grid(True, alpha=0.3)

        axes[1,1].hist(residuals, bins=20, alpha=0.7, edgecolor='black')
        axes[1,1].set_title('Residual Distribution')
        axes[1,1].grid(True, alpha=0.3)
        
        plt.tight_layout(rect=[0, 0, 1, 0.96])
        # Use output_subdir for saving plots
        plot_path = os.path.join(output_subdir, f'plot_{scenario_name}.png')
        plt.savefig(plot_path)
        plt.close(fig)
        doc.add_picture(plot_path, width=Inches(6))

    # Results interpretation from scenario_specific_data
    results_interpretation = scenario_specific_data.get('results_interpretation', {})

    add_heading(doc, 'Results', level=2)
    # Updated to use 'results_narrative' from POST_EXPERIMENT_SCHEMA.json
    add_paragraph(doc, results_interpretation.get('results_narrative', 'N/A'))

    # Stats table (assuming validation_metrics and hypothesis_results are in scenario_specific_data)
    validation_metrics = scenario_specific_data.get('validation_metrics', {})
    hypothesis_results = scenario_specific_data.get('hypothesis_results', {})
    stats_data = {
        "Metric": ["RMSE", "MAE", "MAPE (%)", "NSE", "Pearson's r", "p-value (Pearson)", "Hypothesis Supported"],
        "Value": [
            f"{validation_metrics.get('rmse', 0):.3f}",
            f"{validation_metrics.get('mae', 0):.3f}",
            f"{validation_metrics.get('mape', 0):.2f}",
            f"{validation_metrics.get('nse', 0):.3f}",
            f"{validation_metrics.get('pearson_r', 0):.3f}",
            f"{validation_metrics.get('pearson_p', 0):.4f}",
            str(hypothesis_results.get('hypothesis_supported', 'N/A'))
        ]
    }
    df = pd.DataFrame(stats_data)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Metric']
        row_cells[1].text = row['Value']

    add_heading(doc, 'Discussion', level=2)
    # Updated to use 'discussion_narrative' from POST_EXPERIMENT_SCHEMA.json
    add_paragraph(doc, results_interpretation.get('discussion_narrative', 'N/A'))

    add_heading(doc, 'Conclusion', level=2)
    # Updated to use 'conclusion_narrative' from POST_EXPERIMENT_SCHEMA.json
    add_paragraph(doc, results_interpretation.get('conclusion_narrative', 'N/A'))

    # Save the report
    # Use output_subdir and scenario_name for the report path
    report_path = os.path.join(output_subdir, f'report_{scenario_name}.docx')
    try:
        doc.save(report_path)
        print(f"\nIndividual experiment report for {scenario_name} saved to: {report_path}")
    except Exception as e:
        print(f"Error saving individual report {report_path}: {e}")
    return report_path

def create_comparative_docx_report(llm_response_1, llm_response_2, comparative_output_dir):
    """
    Generates a comparative scientific report in DOCX format from LLM responses.
    llm_response_1: Parsed JSON from the first comparative LLM call (used as fallback).
    llm_response_2: Parsed JSON from the second comparative LLM call (primary source).
    comparative_output_dir: Directory to save the report.
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Main Title
    title = doc.add_heading(level=0)
    title_run = title.add_run('Comparative Analysis of Odum Energy System Model Validation Scenarios')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Default texts
    intro_text = 'Introduction not available.'
    if llm_response_2 and isinstance(llm_response_2, dict):
        intro_text = llm_response_2.get('overall_introduction', intro_text)
    if intro_text == 'Introduction not available.' and llm_response_1 and isinstance(llm_response_1, dict):
        intro_text = llm_response_1.get('overall_introduction', intro_text)

    add_heading(doc, 'Overall Introduction', level=1)
    add_paragraph(doc, intro_text)

    # Experiment Summaries
    add_heading(doc, 'Experiment Summaries', level=1)
    summaries = []
    if llm_response_2 and isinstance(llm_response_2, dict):
        summaries = llm_response_2.get('experiment_summaries', [])

    if not summaries and llm_response_1 and isinstance(llm_response_1, dict): # Fallback if response 2 is empty for summaries
        summaries = llm_response_1.get('experiment_summaries', [])


    for summary_item in summaries:
        if not isinstance(summary_item, dict): continue # Skip if item is not a dict

        scenario_name_text = summary_item.get('scenario_name', 'Unnamed Scenario')
        add_heading(doc, f"Scenario: {scenario_name_text.title()}", level=2)

        add_paragraph(doc, f"Key Finding: {summary_item.get('key_finding', 'N/A')}")

        metrics = summary_item.get('metrics', {})
        if not isinstance(metrics, dict): metrics = {} # Ensure metrics is a dict

        metrics_text = (
            f"RMSE: {metrics.get('rmse', 'N/A')}\n"
            f"NSE: {metrics.get('nse', 'N/A')}\n"
            f"Hypothesis Supported: {metrics.get('hypothesis_supported', 'N/A')}"
        )
        add_paragraph(doc, "Metrics:")
        add_paragraph(doc, metrics_text) # Add metrics as a single block for now

    # Comparative Analysis
    add_heading(doc, 'Comparative Analysis', level=1)
    comparative_analysis_data = {}
    if llm_response_2 and isinstance(llm_response_2, dict):
        comparative_analysis_data = llm_response_2.get('comparative_analysis', {})
    if not isinstance(comparative_analysis_data, dict): comparative_analysis_data = {}


    add_heading(doc, 'Performance Overview', level=2)
    add_paragraph(doc, comparative_analysis_data.get('performance_overview', 'N/A'))

    add_heading(doc, 'Cross-Scenario Insights', level=2)
    add_paragraph(doc, comparative_analysis_data.get('cross_scenario_insights', 'N/A'))

    # Overall Conclusion
    add_heading(doc, 'Overall Conclusion', level=1)
    conclusion_text = 'Overall conclusion not available.'
    if llm_response_2 and isinstance(llm_response_2, dict):
        conclusion_text = llm_response_2.get('overall_conclusion', conclusion_text)

    add_paragraph(doc, conclusion_text)

    # Save Document
    report_path = os.path.join(comparative_output_dir, "comparative_report.docx")
    try:
        doc.save(report_path)
        print(f"\nComparative report successfully saved to: {report_path}")
    except Exception as e:
        print(f"\nError saving comparative report: {e}")

    return report_path